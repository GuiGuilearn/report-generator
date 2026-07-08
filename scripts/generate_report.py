#!/usr/bin/env python3
"""实训报告自动生成系统 — 主控制器。

编排完整流水线：读取 TXT → LLM 提取/生成 → 代码沙箱执行 → 图片渲染 → 模板填充 → 输出 DOCX。

用法:
    python scripts/generate_report.py --date 2026-07-07 [--config config.yaml] [--review]
"""

import argparse
import ast
import json
import logging
import os
import re
import shutil
import sys
import time
import traceback
from dataclasses import dataclass, field
from datetime import datetime

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

import yaml
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from lxml import etree

from scripts.code_sandbox import execute_code_task
from scripts.render_output import render_code_result

# ─── 日志配置 ───
logger = logging.getLogger("generate_report")


# ═══════════════════════════════════════════════════════════════════════
# 数据结构
# ═══════════════════════════════════════════════════════════════════════

@dataclass
class StudentInfo:
    class_name: str = ""
    group: str = ""
    student_id: str = ""
    name: str = ""
    school: str = ""


@dataclass
class TextData:
    course_title: str = ""
    objectives: str = ""
    topics: list[str] = field(default_factory=list)
    details: str = ""
    task_descriptions: list[str] = field(default_factory=list)


@dataclass
class CodeTask:
    name: str = ""
    description: str = ""
    test_inputs: list[str] = field(default_factory=list)
    code: str = ""
    generation_error: str = ""


@dataclass
class ExecutionResult:
    task: CodeTask = field(default_factory=CodeTask)
    status: str = "pending"  # "success" | "failed"
    output_text: str = ""
    image_paths: list[str] = field(default_factory=list)
    error_message: str | None = None


# ═══════════════════════════════════════════════════════════════════════
# System Prompt 模板
# ═══════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT_PHASE1 = """你是一个专业的 Python 实训报告撰写助手。用户会给你一份上课笔记 TXT，
你需要仔细阅读并理解内容，然后输出一个完整的 JSON。

## 输出格式（严格 JSON）

{
    "course_title": "课程标题",
    "objectives": "课程目标（一段话，300字左右）",
    "topics": ["知识点小标题1", "知识点小标题2"],
    "details": "课程内容详情（多段落，500-800字）",
    "task_descriptions": ["任务1的文字描述", "任务2的文字描述"]
}

## 撰写要求

### course_title
- 从笔记中提取本节课的主题

### objectives
- 根据当天所有知识点倒推教学目标
- 用"掌握"、"理解"、"能够"等动词开头
- 覆盖全部知识模块，一段话，语言精炼

### topics
- 提取每个知识模块的名称，只写名称不展开
- 按逻辑顺序排列

### details
- 对每个知识模块展开详细说明
- 用自己的话重新组织，不照抄原文
- 学术化语气，模块间自然过渡
- 总字数 300-600 字

### task_descriptions
- 从笔记中识别需要写代码的练习题和任务
- 每个任务写一段完整的问题描述，包含输入输出要求
- 如果笔记中没有代码练习，返回空数组 []
- 忽略与课程知识无关的操作提示（如注释快捷键、编辑器操作等）
"""

SYSTEM_PROMPT_PHASE2 = """你是一个专业的 Python 代码生成助手。根据任务描述，生成完整可运行的代码。

## 输出格式（严格 JSON）

{
    "name": "任务名称",
    "description": "问题的文字描述",
    "test_inputs": ["模拟输入1", "模拟输入2"],
    "code": "完整可运行代码（含 input 替换器）"
}

## 代码要求
- 语法正确，Python 3 可直接运行
- 包含必要注释，代码风格清晰
- 使用标准库（如 random），禁止使用 os/subprocess/sys 等系统模块
- 严格按任务描述实现具体要求

## 交互式输入处理
如果代码需要 input()，code 必须嵌入以下替换器模板（每个任务独立）：

_test_inputs_{task_index} = ["输入值1", "输入值2", ...]
_idx_{task_index} = 0

def _mock_input_{task_index}(prompt=""):
    global _idx_{task_index}
    print(prompt, end="")
    if _idx_{task_index} < len(_test_inputs_{task_index}):
        val = _test_inputs_{task_index}[_idx_{task_index}]
        _idx_{task_index} += 1
        print(val)
        return val
    return ""

import sys
__builtins__.input = _mock_input_{task_index}
sys.stdin.readline = lambda: _mock_input_{task_index}("")

test_inputs 必须覆盖：正常输入、边界条件、错误输入（汉字、空字符串、非法字符）
替换器之后紧跟原始代码，原始代码不需要修改
代码中禁止重新导入 input 或使用 sys.stdin.readline

纯输出任务：test_inputs 为 []，code 不需要替换器

## 图形输出
如果代码使用 matplotlib 绘图，plt.show() 会被自动拦截保存为图片，
代码正常编写即可，无需额外处理。
"""

SYSTEM_PROMPT_FIX = """你是一个专业的 Python 代码修复助手。以下代码执行时出现了错误，
请仅修复错误，不要改变原有逻辑和功能。

## 原始代码
```
{original_code}
```

## 错误信息
```
{error_message}
```

## 要求
- 仅修复导致错误的代码，不改变其他逻辑
- 保持代码风格和注释不变
- 返回修复后的完整代码（包含 input 替换器，如有）
- 输出格式：严格 JSON，{{ "code": "修复后的完整代码" }}
"""


# ═══════════════════════════════════════════════════════════════════════
# 配置加载
# ═══════════════════════════════════════════════════════════════════════

def load_config(config_path: str) -> dict:
    """加载 YAML 配置文件。

    Args:
        config_path: 配置文件路径

    Returns:
        配置字典

    Raises:
        FileNotFoundError: 配置文件不存在
    """
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"配置文件不存在: {config_path}")
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


# ═══════════════════════════════════════════════════════════════════════
# LLM 服务
# ═══════════════════════════════════════════════════════════════════════

def _parse_json_response(response_text: str) -> dict:
    """解析 LLM 返回的 JSON 响应，处理 markdown 代码块包裹。

    校验步骤：
    1. 去除可能的 markdown 代码块包裹
    2. 尝试 json.loads() 解析
    3. 若失败，返回 None

    Args:
        response_text: LLM 原始响应文本

    Returns:
        解析后的 dict，失败返回 None
    """
    text = response_text.strip()

    # 去除 markdown 代码块包裹
    if text.startswith("```"):
        # 找到第一个换行
        first_nl = text.find("\n")
        if first_nl != -1:
            text = text[first_nl + 1:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                return json.loads(text[start:end + 1])
            except json.JSONDecodeError:
                return None
        return None


def _response_preview(text: str, limit: int = 500) -> str:
    """Return a compact response preview for logs and report diagnostics."""
    preview = " ".join((text or "").split())
    if len(preview) > limit:
        preview = preview[:limit] + "..."
    return preview


def _task_title_from_description(description: str, fallback_index: int) -> str:
    """Build a readable fallback task title from the original task description."""
    text = " ".join((description or "").split())
    if not text:
        return f"任务 {fallback_index + 1}"
    for sep in ["。", "；", ";", "\n"]:
        if sep in text:
            text = text.split(sep, 1)[0]
            break
    return text[:30] + ("..." if len(text) > 30 else "")


def _extract_python_code(response_text: str) -> str:
    """Extract runnable Python code from a plain-code LLM response."""
    text = (response_text or "").strip()
    if not text:
        return ""

    fenced = re.search(r"```(?:python|py)?\s*(.*?)```", text, flags=re.IGNORECASE | re.DOTALL)
    if fenced:
        text = fenced.group(1).strip()

    lines = text.splitlines()
    while lines and not (
        lines[0].lstrip().startswith((
            "import ",
            "from ",
            "def ",
            "class ",
            "#",
            "print(",
            "seats",
            "total",
        ))
        or "=" in lines[0]
    ):
        lines.pop(0)

    while lines and lines[-1].strip() in {"```", "'''", '"""'}:
        lines.pop()

    return "\n".join(lines).strip()


def _is_valid_python(code: str) -> bool:
    """Return True when code is non-empty and can be parsed by Python."""
    if not (code or "").strip():
        return False
    try:
        ast.parse(code)
        return True
    except SyntaxError:
        return False


def _build_railway_ticketing_task(task_description: str) -> CodeTask:
    """Local deterministic fallback for the high-speed railway ticketing exercise."""
    code = r'''import random
import numpy as np

ROWS = 20
COLS = 5
seats = np.zeros((ROWS, COLS), dtype=int)
total_seats = ROWS * COLS


def remaining_seats():
    return int(total_seats - np.sum(seats))


def recommend_seat():
    empty_positions = np.argwhere(seats == 0)
    if len(empty_positions) == 0:
        return None
    row, col = random.choice(empty_positions).tolist()
    return row, col


def parse_seat_pairs(text):
    normalized = text.replace(";", " ").replace(",", " ")
    parts = normalized.split()
    if len(parts) % 2 != 0:
        raise ValueError("请输入成对的行号和列号，例如：0 1 0 2")

    pairs = []
    for index in range(0, len(parts), 2):
        row = int(parts[index])
        col = int(parts[index + 1])
        if row < 0 or row >= ROWS or col < 0 or col >= COLS:
            raise ValueError("座位号超出范围，行号为0-19，列号为0-4")
        pairs.append((row, col))
    return pairs


print("高铁售票程序启动，座位索引从0开始。")
print("输入格式：行号 列号；一次购买多个座位可连续输入，如：0 0 0 1。输入 q 退出。")

while True:
    sold_seats = total_seats - remaining_seats()
    if sold_seats >= total_seats * 0.5:
        seat = recommend_seat()
        if seat is not None:
            print(f"推荐空座：第{seat[0]}排，第{seat[1]}列")

    user_input = input("请输入要购买的座位：").strip()
    if user_input.lower() in {"q", "quit", "exit", "退出"}:
        print("售票结束。")
        break

    try:
        selected_seats = parse_seat_pairs(user_input)
        if not selected_seats:
            print("未输入座位，请重新选择。")
            print(f"当前剩余座位数：{remaining_seats()}")
            continue

        if len(set(selected_seats)) != len(selected_seats):
            print("同一次购买中存在重复座位，无法完成购买。")
            print(f"当前剩余座位数：{remaining_seats()}")
            continue

        unavailable = [(row, col) for row, col in selected_seats if seats[row, col] == 1]
        if unavailable:
            print(f"座位{unavailable}已售，请重新选座，本次购买已回滚。")
            print(f"当前剩余座位数：{remaining_seats()}")
            continue

        for row, col in selected_seats:
            seats[row, col] = 1
        print(f"购买成功，本次购买{len(selected_seats)}个座位：{selected_seats}")
    except ValueError as exc:
        print(f"输入错误：{exc}")

    print(f"当前剩余座位数：{remaining_seats()}")
'''
    return CodeTask(
        name="高铁售票程序",
        description=task_description,
        test_inputs=["0 0 0 1", "0 0", "10 2 10 3", "q"],
        code=code,
    )


def _fallback_code_task_from_description(task_description: str, task_index: int) -> CodeTask | None:
    """Return a local fallback CodeTask for known teaching exercises."""
    text = task_description or ""
    railway_keywords = ["高铁", "售票", "座位", "20", "5"]
    if all(keyword in text for keyword in railway_keywords):
        logger.warning(f"阶段二 [Task {task_index}]: 使用本地高铁售票保底代码。")
        return _build_railway_ticketing_task(task_description)
    return None


def _llm_chat(
    client: OpenAI,
    model: str,
    messages: list[dict],
    temperature: float = 0.3,
    max_tokens: int | None = None,
    timeout: int = 60,
    max_retries: int = 3,
    response_format: dict | None = None,
) -> str:
    """调用 LLM，含重试逻辑。

    Args:
        client: OpenAI 客户端
        model: 模型名称
        messages: 消息列表
        temperature: 温度参数
        max_tokens: 最大 token 数
        timeout: 超时时间（秒）
        max_retries: 最大重试次数
        response_format: 响应格式（如 {"type": "json_object"}）

    Returns:
        LLM 响应文本

    Raises:
        ConnectionError: 网络超时，重试耗尽
    """
    last_error = None
    for attempt in range(max_retries):
        try:
            start = time.time()
            kwargs = {
                "model": model,
                "messages": messages,
                "temperature": temperature,
                "timeout": timeout,
            }
            if max_tokens is not None:
                kwargs["max_tokens"] = max_tokens
            if response_format is not None:
                kwargs["response_format"] = response_format

            response = client.chat.completions.create(**kwargs)
            elapsed = time.time() - start
            logger.info(f"LLM 调用完成，耗时 {elapsed:.1f}s")
            return response.choices[0].message.content
        except Exception as e:
            last_error = e
            if attempt < max_retries - 1:
                wait = 2 * (attempt + 1)
                logger.warning(f"LLM 调用失败 (尝试 {attempt + 1}/{max_retries})，{wait}s 后重试: {e}")
                time.sleep(wait)

    raise ConnectionError(f"LLM 调用失败，已重试 {max_retries} 次: {last_error}")


def ai_generate_text(client: OpenAI, model: str, txt_content: str) -> TextData:
    """阶段一：调用 LLM 提取结构化文本内容。

    Args:
        client: OpenAI 客户端
        model: 模型名称
        txt_content: TXT 笔记内容

    Returns:
        TextData 结构化文本数据

    Raises:
        ValueError: 解析失败或缺少必要字段
    """
    logger.info("阶段一：LLM 提取结构化文本...")
    start = time.time()

    response = _llm_chat(
        client, model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT_PHASE1},
            {"role": "user", "content": txt_content},
        ],
        response_format={"type": "json_object"},
    )

    data = _parse_json_response(response)
    if data is None:
        # 重试一次
        logger.warning("阶段一 JSON 解析失败，重试一次...")
        response = _llm_chat(
            client, model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT_PHASE1},
                {"role": "user", "content": txt_content},
            ],
            response_format={"type": "json_object"},
        )
        data = _parse_json_response(response)

    if data is None:
        raise ValueError(f"阶段一：LLM 返回非标准 JSON，原始响应:\n{response}")

    # 校验必要字段
    required_fields = ["course_title", "objectives", "topics", "details", "task_descriptions"]
    missing = [f for f in required_fields if f not in data]
    if missing:
        raise ValueError(f"阶段一：TextData 缺少必要字段: {missing}")

    elapsed = time.time() - start
    logger.info(f"阶段一完成，耗时 {elapsed:.1f}s，提取到 {len(data.get('topics', []))} 个知识点，"
                f"{len(data.get('task_descriptions', []))} 个代码任务")

    return TextData(
        course_title=data["course_title"],
        objectives=data["objectives"],
        topics=data.get("topics", []),
        details=data["details"],
        task_descriptions=data.get("task_descriptions", []),
    )


def ai_generate_code(
    client: OpenAI,
    model: str,
    task_description: str,
    task_index: int,
) -> CodeTask:
    """阶段二：为单个代码任务调用 LLM 生成代码。

    每个任务独立调用，失败时独立重试，不影响其他任务。

    Args:
        client: OpenAI 客户端
        model: 模型名称
        task_description: 任务描述文本
        task_index: 任务索引（用于生成独立变量名）

    Returns:
        CodeTask 对象

    Raises:
        ValueError: 解析失败或缺少必要字段
    """
    logger.info(f"阶段二 [Task {task_index}]: 生成代码...")

    # 将 {task_index} 替换为实际索引
    prompt = SYSTEM_PROMPT_PHASE2.replace("{task_index}", str(task_index))

    response = _llm_chat(
        client, model,
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": task_description},
        ],
        response_format={"type": "json_object"},
        max_tokens=4000,
    )

    data = _parse_json_response(response)
    if data is None:
        # 重试一次
        logger.warning(f"阶段二 [Task {task_index}]: JSON 解析失败，重试一次...")
        response = _llm_chat(
            client, model,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": task_description},
            ],
            response_format={"type": "json_object"},
            max_tokens=4000,
        )
        data = _parse_json_response(response)

    if data is None:
        fallback_task = _fallback_code_task_from_description(task_description, task_index)
        if fallback_task is not None:
            return fallback_task

        logger.warning(f"阶段二 [Task {task_index}]: JSON 解析仍失败，改用纯 Python 代码生成...")
        plain_response = _llm_chat(
            client, model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是 Python 代码生成助手。只返回完整 Python 3 源代码，"
                        "不要返回 JSON，不要解释，不要使用 Markdown。"
                    ),
                },
                {"role": "user", "content": task_description},
            ],
            max_tokens=4000,
            response_format=None,
        )
        plain_code = _extract_python_code(plain_response)
        if _is_valid_python(plain_code):
            logger.warning(f"阶段二 [Task {task_index}]: 已从纯代码响应中恢复源代码。")
            return CodeTask(
                name=_task_title_from_description(task_description, task_index),
                description=task_description,
                test_inputs=["q"],
                code=plain_code,
            )

        raise ValueError(
            "LLM 返回非标准 JSON，纯代码恢复也失败。"
            f"\n任务描述：{task_description}"
            f"\nJSON 响应摘要：{_response_preview(response)}"
            f"\n纯代码响应摘要：{_response_preview(plain_response)}"
        )

    required_fields = ["name", "description", "test_inputs", "code"]
    missing = [f for f in required_fields if f not in data]
    if missing:
        fallback_task = _fallback_code_task_from_description(task_description, task_index)
        if fallback_task is not None:
            return fallback_task
        raise ValueError(f"阶段二 [Task {task_index}]: CodeTask 缺少必要字段: {missing}")

    if not _is_valid_python(data.get("code", "")):
        fallback_task = _fallback_code_task_from_description(task_description, task_index)
        if fallback_task is not None:
            return fallback_task
        raise ValueError(
            f"阶段二 [Task {task_index}]: LLM 返回的 code 字段为空或不是合法 Python。"
            f"\n任务描述：{task_description}"
            f"\n原始响应摘要：{_response_preview(response)}"
        )

    return CodeTask(
        name=data["name"],
        description=data["description"],
        test_inputs=data.get("test_inputs", []),
        code=data["code"],
    )


def ai_fix_code(client: OpenAI, model: str, original_code: str, error_message: str) -> str | None:
    """调用 LLM 修复代码错误。

    Args:
        client: OpenAI 客户端
        model: 模型名称
        original_code: 原始代码
        error_message: 错误信息

    Returns:
        修复后的代码，失败返回 None
    """
    prompt = SYSTEM_PROMPT_FIX.format(
        original_code=original_code,
        error_message=error_message,
    )

    try:
        response = _llm_chat(
            client, model,
            messages=[
                {"role": "system", "content": prompt},
            ],
            response_format={"type": "json_object"},
            max_retries=1,
        )
        data = _parse_json_response(response)
        if data and "code" in data:
            return data["code"]
        return None
    except Exception as e:
        logger.warning(f"自动修复 LLM 调用失败: {e}")
        return None


# ═══════════════════════════════════════════════════════════════════════
# 模板引擎
# ═══════════════════════════════════════════════════════════════════════

# Word 文档 XML 命名空间
WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _find_bookmark_paragraph(doc: Document, bookmark_name: str):
    """在文档中查找指定书签所在的段落。

    Args:
        doc: Document 对象
        bookmark_name: 书签名称

    Returns:
        (paragraph, parent_element) 元组，未找到返回 (None, None)
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for elem in run._element.iter():
                if elem.tag == f"{{{WML_NS}}}bookmarkStart":
                    name = elem.get(f"{{{WML_NS}}}name")
                    if name == bookmark_name:
                        return paragraph, paragraph._element.getparent()
    return None, None


def _insert_paragraph_after(paragraph, text: str, style: str | None = None) -> Paragraph:
    """在指定段落后插入新段落。

    Args:
        paragraph: 参考段落
        text: 新段落文本
        style: 段落样式名
    """
    parent = paragraph._element.getparent()
    index = list(parent).index(paragraph._element)

    new_p = etree.SubElement(parent, f"{{{WML_NS}}}p", {})
    # 插入到参考段落之后
    parent.remove(new_p)
    parent.insert(index + 1, new_p)

    if style:
        pPr = etree.SubElement(new_p, f"{{{WML_NS}}}pPr")
        pStyle = etree.SubElement(pPr, f"{{{WML_NS}}}pStyle")
        pStyle.set(f"{{{WML_NS}}}val", style)

    run_elem = etree.SubElement(new_p, f"{{{WML_NS}}}r")
    t_elem = etree.SubElement(run_elem, f"{{{WML_NS}}}t")
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return Paragraph(new_p, paragraph._parent)


def fill_at_bookmark(
    doc: Document,
    bookmark_name: str,
    content_insert_fn,
    fallback_markers: list[str] | None = None,
) -> None:
    """在指定书签位置插入内容。

    通过书签名称定位插入点，调用 content_insert_fn 在书签段落后插入内容。

    Args:
        doc: Document 对象
        bookmark_name: 书签名称
        content_insert_fn: 内容插入回调函数，签名为 fn(paragraph, parent_element)
    """
    para, parent = _find_bookmark_paragraph(doc, bookmark_name)
    if para is None:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if fallback_markers and any(marker in text for marker in fallback_markers):
                para = paragraph
                parent = paragraph._element.getparent()
                logger.info(f"未找到书签 {bookmark_name}，改用标题定位: {text}")
                break
        else:
            logger.warning(f"未找到书签或标题: {bookmark_name}")
            return
    content_insert_fn(para, parent)


def _fill_cover_info(doc: Document, student: StudentInfo) -> None:
    """填充封面信息（替换文档中的占位符）。

    在文档所有段落中查找 {key} 形式的占位符并替换为 config 中的值。

    Args:
        doc: Document 对象
        student: 学生信息
    """
    field_map = {
        "{class_name}": student.class_name,
        "{group}": student.group,
        "{student_id}": student.student_id,
        "{name}": student.name,
        "{school}": student.school,
    }

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in field_map.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

        text = paragraph.text.strip()
        label_map = {
            "学    院：": student.school,
            "专业班级：": student.class_name,
            "学    号：": student.student_id,
            "姓    名：": student.name,
        }
        for label, value in label_map.items():
            if text.startswith(label) and value not in text:
                paragraph.text = f"{label} {value}"


def _fill_objectives(doc: Document, objectives: str) -> None:
    """在 bookmark_objectives 处插入课程目标。

    Args:
        doc: Document 对象
        objectives: 课程目标文本
    """
    def insert_fn(para, parent):
        _insert_paragraph_after(para, objectives)
    fill_at_bookmark(doc, "bookmark_objectives", insert_fn, ["课程目标"])


def _fill_topics(doc: Document, topics: list[str]) -> None:
    """在 bookmark_topics 处插入知识点列表（逐条编号）。

    Args:
        doc: Document 对象
        topics: 知识点列表
    """
    def insert_fn(para, parent):
        for i, topic in enumerate(topics, 1):
            para = _insert_paragraph_after(para, f"{i}. {topic}")
            # 更新参考段落为最新插入的段落
    fill_at_bookmark(doc, "bookmark_topics", insert_fn, ["课程内容："])


def _fill_details(doc: Document, details: str) -> None:
    """在 bookmark_details 处插入课程内容详情。

    支持多段落（按 \\n\\n 分割为段落）。

    Args:
        doc: Document 对象
        details: 课程详情文本
    """
    def insert_fn(para, parent):
        paragraphs = [p.strip() for p in details.split("\n\n") if p.strip()]
        if not paragraphs:
            paragraphs = [details]
        for text in paragraphs:
            para = _insert_paragraph_after(para, text)
    fill_at_bookmark(doc, "bookmark_details", insert_fn, ["课程内容详情"])


def _insert_image_after(doc: Document, paragraph, image_path: str, width_inches: float = 6.3) -> object:
    """在指定段落后插入一个居中的图片段落。

    使用 python-docx API 创建图片段落，然后将其 XML 元素移动到书签位置之后。

    Args:
        doc: python-docx Document 对象
        paragraph: 参考段落
        image_path: 图片文件路径
        width_inches: 图片宽度（英寸）

    Returns:
        新插入的图片段落对象
    """
    parent = paragraph._element.getparent()
    index = list(parent).index(paragraph._element)

    # 使用 python-docx 创建图片段落（会添加到文档末尾）
    tmp_para = doc.add_paragraph()
    tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tmp_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    # 将段落元素从末尾移动到书签位置之后
    tmp_para._element.getparent().remove(tmp_para._element)
    parent.insert(index + 1, tmp_para._element)

    return tmp_para


def _fill_code_tasks(
    doc: Document,
    results: list[ExecutionResult],
    day: str,
) -> None:
    """在 bookmark_code 处插入代码任务及运行结果。

    每个任务插入：
    - 任务名称（标题）
    - 成功：渲染后的 PNG 图片 + 题注
    - 失败：错误信息 + 源代码

    Args:
        doc: Document 对象
        results: 执行结果列表
        day: 日期字符串（用于题注编号）
    """
    def insert_fn(para, parent):
        for i, result in enumerate(results, 1):
            task = result.task

            # 任务标题
            para = _insert_paragraph_after(para, f"任务 {i}：{task.name}")
            if task.description:
                para = _insert_paragraph_after(para, f"任务描述：{task.description}")

            para = _insert_paragraph_after(para, "源代码：")
            if task.code:
                para = _insert_paragraph_after(para, task.code)
            else:
                para = _insert_paragraph_after(para, "未生成源代码。请根据上方任务描述和下方失败原因手动补充。")

            if result.status == "success":
                if result.image_paths:
                    para = _insert_paragraph_after(para, "运行结果：")
                    # 插入图片
                    for img_path in result.image_paths:
                        if os.path.exists(img_path):
                            # 插入图片段落
                            img_para = _insert_image_after(doc, para, img_path)
                            para._element = img_para._element

                            # 题注
                            caption = f"图 {day}-{i}  {task.name}"
                            para = _insert_paragraph_after(para, caption)
                else:
                    # 成功但无图片（如渲染失败），展示输出文本
                    para = _insert_paragraph_after(para, "状态：执行成功（图片渲染失败）")
                    if result.output_text:
                        para = _insert_paragraph_after(para, f"运行输出：\n{result.output_text[:500]}")
            else:
                # 失败状态
                para = _insert_paragraph_after(para, "状态：代码执行失败")

                error_msg = result.error_message or "未知错误"
                para = _insert_paragraph_after(para, f"错误信息：{error_msg}")

    fill_at_bookmark(doc, "bookmark_code", insert_fn, ["课程代码及执行结果"])


def _remove_template_instructions(doc: Document) -> None:
    """Remove authoring instructions left in the Word template."""
    markers = (
        "要求如下",
        "1）代码必须是文字",
        "2）结果必须是图片",
        "3）正文字体",
        "英文及数字字体",
        "4）报告命名",
        "(严格按照此格式命名)",
    )
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if any(marker in text for marker in markers):
            paragraph._element.getparent().remove(paragraph._element)


# ═══════════════════════════════════════════════════════════════════════
# 主流水线
# ═══════════════════════════════════════════════════════════════════════

def generate_report(
    date_str: str,
    config_path: str = "config.yaml",
    model_override: str | None = None,
    provider_override: str | None = None,
    api_key_override: str | None = None,
    base_url_override: str | None = None,
    review: bool = False,
) -> int:
    """主流水线：从 TXT 到 DOCX 的端到端自动化。

    Args:
        date_str: 日期字符串，格式 YYYY-MM-DD
        config_path: 配置文件路径
        model_override: 覆盖配置文件中的模型名称
        review: 是否启用 review 模式

    Returns:
        退出码：0 成功，1 失败
    """
    # Step 0: 加载配置
    logger.info(f"===== 实训报告生成开始 =====")
    logger.info(f"日期: {date_str}")

    config = load_config(config_path)
    llm_config = config.get("llm", {})
    provider_name = provider_override or llm_config.get("provider")
    provider_config = {}
    if provider_name:
        provider_config = llm_config.get("providers", {}).get(provider_name, {})

    model = model_override or provider_config.get("model") or llm_config.get("model")
    api_key = api_key_override or provider_config.get("api_key") or llm_config.get("api_key")
    api_key_env = provider_config.get("api_key_env") or llm_config.get("api_key_env") or "OPENAI_API_KEY"
    if not api_key:
        api_key = os.environ.get(api_key_env) or os.environ.get("OPENAI_API_KEY")
    base_url = (
        base_url_override
        or provider_config.get("base_url")
        or llm_config.get("base_url")
        or os.environ.get("OPENAI_BASE_URL")
    )

    if not api_key:
        logger.error(
            "缺少 API Key。请使用 --api-key，或设置环境变量 %s，或在 config.yaml 的 llm/providers 中配置 api_key/api_key_env。",
            api_key_env,
        )
        return 1
    api_key = api_key.strip()
    if api_key in {"你的DeepSeekKey", "你的OpenAIKey", "sk-xxx", "your-api-key"}:
        logger.error("API Key 仍是示例占位符，请替换为真实 key。当前读取变量: %s", api_key_env)
        return 1
    try:
        api_key.encode("ascii")
    except UnicodeEncodeError:
        logger.error(
            "API Key 包含非 ASCII 字符，不能包含中文或全角字符。请重新设置真实 key。当前读取变量: %s",
            api_key_env,
        )
        return 1
    if not model:
        logger.error("缺少模型名称。请使用 --model，或在 config.yaml 的 llm/providers 中配置 model。")
        return 1

    logger.info(f"LLM provider: {provider_name or 'default'}")
    logger.info(f"LLM model: {model}")
    if base_url:
        logger.info(f"LLM base_url: {base_url}")
    timeout = config["execution"]["timeout_per_cell"]
    auto_fix_max = config["execution"]["auto_fix_max_attempts"]

    # 清空 temp/ 目录
    temp_dir = "temp/"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir, exist_ok=True)

    # 初始化 LLM 客户端
    client = OpenAI(
        api_key=api_key,
        base_url=base_url,
    )

    # 确保输出目录存在
    os.makedirs("outputs", exist_ok=True)
    os.makedirs("logs", exist_ok=True)

    all_success = True
    text_data = None
    code_tasks = []
    execution_results = []

    try:
        # Step 1: 读取 TXT 文件
        txt_path = f"daily_notes/{date_str}.txt"
        if not os.path.exists(txt_path):
            logger.error(f"TXT 文件不存在: {txt_path}")
            return 1

        with open(txt_path, "r", encoding="utf-8") as f:
            txt_content = f.read()

        # TXT 预处理：rstrip 清理尾部空白
        txt_content = txt_content.rstrip()
        logger.info(f"TXT 文件读取完成，{len(txt_content)} 字符")

        # Step 2: 阶段一 — LLM 提取结构化文本
        text_data = ai_generate_text(client, model, txt_content)
        logger.info(f"课程标题: {text_data.course_title}")
        logger.info(f"知识点: {len(text_data.topics)} 个")
        logger.info(f"代码任务描述: {len(text_data.task_descriptions)} 个")

        # Step 3: 阶段二 — 逐一生成代码
        if text_data.task_descriptions:
            for i, desc in enumerate(text_data.task_descriptions):
                try:
                    task = ai_generate_code(client, model, desc, i)
                    code_tasks.append(task)
                    logger.info(f"  [Task {i}] {task.name} — 代码生成完成")
                except Exception as e:
                    logger.error(f"  [Task {i}] 代码生成失败: {e}")
                    # 创建一个失败的 CodeTask
                    code_tasks.append(CodeTask(
                        name=_task_title_from_description(desc, i),
                        description=desc,
                        code="",
                        generation_error=str(e),
                    ))
        else:
            logger.warning("没有代码任务需要生成")

        # Step 4: 代码沙箱 — 执行每个任务
        for i, task in enumerate(code_tasks):
            logger.info(f"代码沙箱 [Task {i}]: {task.name}")

            if not task.code:
                error_message = (
                    "代码生成失败，未能得到可运行代码。\n"
                    f"任务描述：{task.description or task.name}"
                )
                if task.generation_error:
                    error_message += f"\n失败原因：{task.generation_error}"
                execution_results.append(ExecutionResult(
                    task=task,
                    status="failed",
                    error_message=error_message,
                ))
                all_success = False
                continue

            # 创建自动修复回调
            def make_fix_fn(c, m):
                return lambda code, err: ai_fix_code(c, m, code, err)

            fix_fn = make_fix_fn(client, model)

            success, stdout, stderr, final_code = execute_code_task(
                code=task.code,
                test_inputs=task.test_inputs,
                task_index=i,
                timeout=timeout,
                auto_fix_fn=fix_fn,
                auto_fix_max_attempts=auto_fix_max,
            )

            # 更新 task.code 为最终代码（可能经过修复）
            task.code = final_code

            # 检查图形输出
            figure_path = os.path.join("temp", "_output_figure.png")
            has_figure = os.path.exists(figure_path)

            # 渲染图片
            image_paths = []
            if success:
                output_text = stdout
                if stderr:
                    output_text += f"\n[stderr]\n{stderr}"

                img_save_path = f"temp/task_{i:02d}_{task.name}.png"

                # 确定图形输出路径
                extra_img = figure_path if has_figure else None
                try:
                    render_code_result(
                        task_name=task.name,
                        code=task.code,
                        output_text=output_text,
                        extra_image_path=extra_img,
                        save_path=img_save_path,
                    )
                    image_paths.append(img_save_path)
                except Exception as e:
                    logger.warning(f"[Task {i}] 图片渲染失败: {e}")

                execution_results.append(ExecutionResult(
                    task=task,
                    status="success",
                    output_text=output_text,
                    image_paths=image_paths,
                ))
                logger.info(f"  [Task {i}] OK 执行成功")
            else:
                error_msg = stderr or stdout or "未知错误"
                execution_results.append(ExecutionResult(
                    task=task,
                    status="failed",
                    output_text=stdout,
                    error_message=error_msg,
                ))
                all_success = False
                logger.warning(f"  [Task {i}] FAIL 执行失败: {error_msg[:100]}")

            # 清理图形输出临时文件
            if has_figure:
                try:
                    os.remove(figure_path)
                except Exception:
                    pass

        # Step 4 总结
        success_count = sum(1 for r in execution_results if r.status == "success")
        total_count = len(execution_results)
        logger.info(f"代码任务执行完成: {success_count}/{total_count} 成功")

        # Step 5: 图片渲染（已在 Step 4 中完成）

        # Step 6: 模板填充
        template_path = "template/实训报告模板.docx"
        if not os.path.exists(template_path):
            logger.error(f"模板文件不存在: {template_path}")
            return 1

        doc = Document(template_path)

        # 填充封面信息
        student = StudentInfo(
            class_name=config["student"]["class_name"],
            group=config["student"]["group"],
            student_id=config["student"]["student_id"],
            name=config["student"]["name"],
            school=config["student"]["school"],
        )
        _fill_cover_info(doc, student)

        # 填充四个章节
        if text_data:
            _fill_objectives(doc, text_data.objectives)
            _fill_topics(doc, text_data.topics)
            _fill_details(doc, text_data.details)
            _fill_code_tasks(doc, execution_results, date_str)
            _remove_template_instructions(doc)

        # 输出文件命名
        naming = config["output"]["naming"]
        output_filename = naming.format(
            class_name=student.class_name,
            group=student.group,
            student_id=student.student_id,
            name=student.name,
        )
        output_path = f"outputs/{output_filename}"
        try:
            doc.save(output_path)
        except PermissionError:
            stem, ext = os.path.splitext(output_filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"outputs/{stem}_{timestamp}{ext}"
            logger.warning(f"输出文件被占用，已改为另存: {output_path}")
            doc.save(output_path)
        logger.info(f"报告已生成: {output_path}")

        # Step 7: 写入日志
        log_path = f"logs/{date_str}.log"
        # 日志已通过 logging 配置处理

        # --review 模式
        if review:
            print()
            print("=" * 40)
            print("  实训报告生成摘要")
            print("=" * 40)
            print(f"课程标题：{text_data.course_title if text_data else 'N/A'}")
            print(f"课程目标：{text_data.objectives[:50]}..." if text_data and text_data.objectives else "N/A")
            print(f"知识点：{len(text_data.topics) if text_data else 0} 个")
            print(f"代码任务：{total_count} 个")
            for r in execution_results:
                status = "OK" if r.status == "success" else "FAIL"
                msg = r.task.name
                if r.error_message:
                    msg += f" — {r.error_message[:80]}"
                print(f"  {status} {msg}")
            print(f"成功率：{success_count}/{total_count}")
            print("=" * 40)

            answer = input("是否确认生成报告？[Y/n] ").strip().lower()
            if answer and answer != "y":
                logger.info("用户取消生成")
                return 0

        # 确定退出码
        if not text_data:
            return 1
        if not execution_results:
            return 0
        if success_count == 0 and total_count > 0:
            return 1

        return 0

    except Exception as e:
        logger.error(f"流水线异常: {e}")
        logger.error(traceback.format_exc())
        return 1

    finally:
        # Step 8: 清理 temp/ 临时文件
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                logger.info("临时目录已清理")
            except Exception as e:
                logger.warning(f"清理临时目录失败: {e}")


# ═══════════════════════════════════════════════════════════════════════
# CLI 入口
# ═══════════════════════════════════════════════════════════════════════

def setup_logging(date_str: str) -> None:
    """配置日志：同时输出到控制台和文件。

    Args:
        date_str: 日期字符串，用于日志文件名
    """
    os.makedirs("logs", exist_ok=True)
    log_path = f"logs/{date_str}.log"

    root_logger = logging.getLogger("generate_report")
    root_logger.setLevel(logging.DEBUG)

    # 控制台 handler
    console = logging.StreamHandler(sys.stdout)
    console.setLevel(logging.INFO)
    console_fmt = logging.Formatter(
        "[%(asctime)s] %(levelname)s: %(message)s",
        datefmt="%H:%M:%S",
    )
    console.setFormatter(console_fmt)

    # 文件 handler
    file_handler = logging.FileHandler(log_path, encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_fmt = logging.Formatter(
        "%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )
    file_handler.setFormatter(file_fmt)

    root_logger.addHandler(console)
    root_logger.addHandler(file_handler)

    # 同时配置 code_sandbox 和 render_output 的日志
    for mod_name in ["scripts.code_sandbox", "scripts.render_output"]:
        mod_logger = logging.getLogger(mod_name)
        mod_logger.setLevel(logging.DEBUG)
        mod_logger.addHandler(console)
        mod_logger.addHandler(file_handler)


def main():
    parser = argparse.ArgumentParser(
        description="实训报告自动生成系统",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    python scripts/generate_report.py --date 2026-07-07
    python scripts/generate_report.py --date 2026-07-07 --config my_config.yaml
    python scripts/generate_report.py --date 2026-07-07 --review
        """,
    )
    parser.add_argument(
        "--date", required=True,
        help="日期，格式 YYYY-MM-DD（对应 daily_notes/ 下的 TXT 文件名）",
    )
    parser.add_argument(
        "--config", default="config.yaml",
        help="配置文件路径（默认: config.yaml）",
    )
    parser.add_argument(
        "--model", default=None,
        help="覆盖 LLM 模型（默认来自 config.yaml）",
    )
    parser.add_argument(
        "--provider", default=None,
        help="选择 config.yaml 中 llm.providers 下的服务商配置，例如 deepseek/openai",
    )
    parser.add_argument(
        "--api-key", default=None,
        help="临时指定 API Key；优先级高于环境变量和 config.yaml",
    )
    parser.add_argument(
        "--base-url", default=None,
        help="临时指定 OpenAI-compatible API Base URL，例如 https://api.deepseek.com",
    )
    parser.add_argument(
        "--review", action="store_true",
        help="启用 review 模式：流程结束打印摘要，等待用户确认",
    )

    args = parser.parse_args()

    # 设置日志
    setup_logging(args.date)

    # 运行流水线
    exit_code = generate_report(
        date_str=args.date,
        config_path=args.config,
        model_override=args.model,
        provider_override=args.provider,
        api_key_override=args.api_key,
        base_url_override=args.base_url,
        review=args.review,
    )

    sys.exit(exit_code)


if __name__ == "__main__":
    main()
